"""DOCX conversion functionality"""

import re
import markdown
from docx import Document
from docx.shared import Pt, Inches
from bs4 import BeautifulSoup
from io import BytesIO
from .text_formatter import add_formatted_text
from .mermaid_docx_handler import extract_mermaid_and_render
from .page_break_handler import convert_page_breaks_to_placeholder
from .obsidian_to_html import preprocess_obsidian_for_html


def convert_to_docx(markdown_content, render_mermaid=True, obsidian_mode=True):
    """Convert Markdown content to DOCX format

    Args:
        markdown_content: The markdown content to convert
        render_mermaid: Whether to render Mermaid diagrams (default: True)
        obsidian_mode: Whether to preprocess Obsidian syntax (default: True)

    Returns:
        BytesIO buffer containing the DOCX document
    """
    # Preprocess Obsidian syntax if enabled (like PDF converter)
    if obsidian_mode:
        # Note: We use HTML-specific preprocessing for DOCX output
        # (unlike PDF which uses LaTeX-based preprocessing)
        # This ensures callouts, checkboxes, etc. render properly in Word
        markdown_content = preprocess_obsidian_for_html(markdown_content)

    # Handle Mermaid diagrams if enabled
    diagram_images = {}
    if render_mermaid:
        try:
            markdown_content, diagram_images = extract_mermaid_and_render(markdown_content)
            if diagram_images:
                print(f"Rendered {len(diagram_images)} Mermaid diagrams")
        except Exception as e:
            print(f"Warning: Mermaid rendering disabled due to error: {e}")

    # Preprocess page break markers before HTML conversion
    # Using shared module for consistency with PDF converter
    markdown_content = convert_page_breaks_to_placeholder(markdown_content)

    # Convert markdown to HTML first
    html_content = markdown.markdown(
        markdown_content,
        extensions=['extra', 'codehilite', 'tables', 'fenced_code']
    )

    print(f"DEBUG: HTML content snippet: {html_content[:500]}...")

    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # Create a new Word document
    doc = Document()

    # Track the previous element type to detect list boundaries
    prev_list_type = None

    # Process HTML elements
    for element in soup.children:
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Add heading with formatting
            level = int(element.name[1])
            para = doc.add_heading(level=level)
            add_formatted_text(para, element)
            # Reset list tracking after heading
            prev_list_type = None
        elif element.name == 'p':
            # Check if this paragraph contains a page break marker
            text = element.get_text().strip()
            if text == '|||PAGEBREAK|||':
                # Insert page break
                doc.add_page_break()
                # Reset list tracking after page break
                prev_list_type = None
                continue

            # Check if this paragraph contains a Mermaid placeholder
            # Get all text including from child elements like <strong>
            match = re.search(r'MERMAID_DIAGRAM_(\d+)', text)  # Match without underscores too
            if match and diagram_images:
                # This is a Mermaid placeholder - insert image directly
                diagram_id = match.group(1)
                image_key = f"mermaid_{diagram_id}"
                if image_key in diagram_images:
                    para = doc.add_paragraph()
                    try:
                        image_buffer = diagram_images[image_key]
                        image_buffer.seek(0)
                        run = para.add_run()
                        run.add_picture(image_buffer, width=Inches(5.0))
                        print(f"✓ Inserted Mermaid diagram {diagram_id} during HTML processing")
                    except Exception as e:
                        para.add_run(f"[Error: Could not insert Mermaid diagram: {e}]")
                        print(f"Error inserting diagram {diagram_id}: {e}")
                else:
                    print(f"Warning: Image key {image_key} not found in diagram_images")
            else:
                # Regular paragraph with formatting
                para = doc.add_paragraph()
                add_formatted_text(para, element)
            # Reset list tracking after paragraph
            prev_list_type = None
        elif element.name in ['ul', 'ol']:
            # Check if we're continuing from a previous list of same type
            current_list_type = element.name

            # If starting a new list after a break, add empty paragraph to reset numbering
            if prev_list_type is None and current_list_type == 'ol':
                # This forces Word to restart numbering
                pass

            # Add list items with formatting
            for li in element.find_all('li', recursive=False):
                # Check if this is a task list item (checkbox)
                li_text = li.get_text()
                checkbox_match = re.match(r'^\s*\[([ xX])\]\s*(.*)$', li_text, re.DOTALL)

                if checkbox_match:
                    # This is a checkbox item
                    checked = checkbox_match.group(1).lower() == 'x'
                    text_content = checkbox_match.group(2)

                    # Use plain paragraph style (no bullet)
                    para = doc.add_paragraph()
                    # Add checkbox symbol
                    checkbox_symbol = '☑' if checked else '☐'
                    run = para.add_run(f"{checkbox_symbol} ")
                    # Add the rest of the text
                    para.add_run(text_content)
                else:
                    # Regular list item
                    para = doc.add_paragraph(style='List Bullet' if element.name == 'ul' else 'List Number')
                    add_formatted_text(para, li)

            # Update previous list type
            prev_list_type = current_list_type
        elif element.name == 'table':
            # Add table
            rows = element.find_all('tr')
            if rows:
                # Count columns from first row
                first_row = rows[0]
                cols = len(first_row.find_all(['th', 'td']))

                # Create table in document
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'

                # Fill table data
                for row_idx, row in enumerate(rows):
                    cells = row.find_all(['th', 'td'])
                    for col_idx, cell in enumerate(cells):
                        table_cell = table.rows[row_idx].cells[col_idx]
                        # Clear default text
                        table_cell.text = ''
                        # Add formatted text
                        if table_cell.paragraphs:
                            para = table_cell.paragraphs[0]
                        else:
                            para = table_cell.add_paragraph()
                        add_formatted_text(para, cell)

                        # Bold header cells
                        if cell.name == 'th':
                            for paragraph in table_cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
        elif element.name in ['pre', 'code']:
            # Add code block
            para = doc.add_paragraph(element.get_text())
            para.style = 'No Spacing'
            for run in para.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
