"""Inline element formatting for Word documents"""

from docx.shared import Pt
from .hyperlink import add_hyperlink


def process_inline_element(paragraph, elem):
    """Process inline elements recursively with formatting"""
    if elem.name == 'a':
        # Handle hyperlinks
        url = elem.get('href', '')
        if url:
            add_hyperlink(paragraph, url, elem.get_text())
    elif elem.name in ['strong', 'b']:
        # Bold text - process children
        for child in elem.children:
            if isinstance(child, str):
                run = paragraph.add_run(child)
                run.bold = True
            else:
                # Nested formatting (e.g., bold + italic)
                for text in child.stripped_strings:
                    run = paragraph.add_run(text)
                    run.bold = True
                    if child.name in ['em', 'i']:
                        run.italic = True
    elif elem.name in ['em', 'i']:
        # Italic text - process children
        for child in elem.children:
            if isinstance(child, str):
                run = paragraph.add_run(child)
                run.italic = True
            else:
                for text in child.stripped_strings:
                    run = paragraph.add_run(text)
                    run.italic = True
                    if child.name in ['strong', 'b']:
                        run.bold = True
    elif elem.name == 'code':
        # Inline code
        run = paragraph.add_run(elem.get_text())
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
    else:
        # Regular text
        paragraph.add_run(elem.get_text())
