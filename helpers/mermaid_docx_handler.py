"""Handle Mermaid diagrams in DOCX conversion"""

import re
from docx.shared import Inches
from .mermaid_detector import detect_mermaid_blocks
from .mermaid_renderer import render_mermaid_to_png


def extract_mermaid_and_render(markdown_content):
    """
    Extract Mermaid diagrams and render them to images

    Returns:
        modified_markdown: Markdown with Mermaid blocks replaced with placeholders
        diagram_images: Dict mapping placeholder IDs to image buffers
    """
    mermaid_blocks = detect_mermaid_blocks(markdown_content)

    print(f"DEBUG: Found {len(mermaid_blocks)} Mermaid blocks")

    if not mermaid_blocks:
        return markdown_content, {}

    diagram_images = {}
    modified_content = markdown_content

    # Process blocks in reverse to maintain positions
    for idx, (start, end, code) in enumerate(reversed(mermaid_blocks)):
        diagram_id = len(mermaid_blocks) - idx - 1

        try:
            print(f"DEBUG: Rendering Mermaid diagram {diagram_id}...")
            # Render to PNG
            image_buffer = render_mermaid_to_png(code)
            diagram_images[f"mermaid_{diagram_id}"] = image_buffer

            # Replace with placeholder
            placeholder = f"__MERMAID_DIAGRAM_{diagram_id}__"
            modified_content = modified_content[:start] + placeholder + modified_content[end:]
            print(f"DEBUG: Replaced diagram {diagram_id} with placeholder: {placeholder}")

        except Exception as e:
            # If rendering fails, keep the code block
            print(f"Warning: Failed to render Mermaid diagram {diagram_id}: {e}")

    print(f"DEBUG: Modified content includes: {modified_content[:200]}...")
    return modified_content, diagram_images


def insert_mermaid_images_in_doc(doc, diagram_images):
    """
    Find placeholder text in document and replace with images

    Args:
        doc: python-docx Document object
        diagram_images: Dict mapping diagram IDs to image buffers
    """
    # Check all paragraphs including those in tables
    paragraphs_to_check = list(doc.paragraphs)

    # Also check paragraphs in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs_to_check.extend(cell.paragraphs)

    for paragraph in paragraphs_to_check:
        text = paragraph.text.strip()

        # Check if this paragraph contains a mermaid placeholder
        match = re.search(r'__MERMAID_DIAGRAM_(\d+)__', text)
        if match:
            diagram_id = match.group(1)
            image_key = f"mermaid_{diagram_id}"

            if image_key in diagram_images:
                # Clear the placeholder text
                paragraph.clear()

                # Add the image
                run = paragraph.add_run()
                try:
                    image_buffer = diagram_images[image_key]
                    image_buffer.seek(0)
                    run.add_picture(image_buffer, width=Inches(5.0))
                    print(f"Successfully inserted Mermaid diagram {diagram_id}")
                except Exception as e:
                    # If image insertion fails, add error text
                    paragraph.add_run(f"[Error: Could not insert Mermaid diagram: {e}]")
                    print(f"Error inserting diagram {diagram_id}: {e}")
